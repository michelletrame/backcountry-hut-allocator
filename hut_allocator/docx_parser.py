"""
DOCX Parser for extracting reservation requests from Word documents.
"""
from docx import Document
import re
from datetime import datetime

class DOCXParser:
    """Parse reservation request data from Word documents."""

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = None
        self.data = {
            'leader_name': None,
            'email': None,
            'phone': None,
            'preferences': []
        }

    def parse(self):
        """
        Parse the DOCX and extract structured data.
        Returns dict with leader info and preferences.
        """
        self.doc = Document(self.docx_path)

        # Extract from tables (most common for forms)
        self._extract_from_tables()

        # If tables didn't work, try paragraph parsing
        if not self.data['preferences']:
            self._extract_from_paragraphs()

        return self.data

    def _extract_from_tables(self):
        """Extract data from Word tables."""
        for table in self.doc.tables:
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]

                # Look for leader information
                if len(cells) >= 2:
                    if 'Leader Name' in cells[0]:
                        self.data['leader_name'] = cells[1]
                    elif 'Email' in cells[0]:
                        self.data['email'] = cells[1]
                    elif 'Phone' in cells[0]:
                        self.data['phone'] = cells[1]

                    # Look for hut preferences
                    elif 'Hut Preference' in cells[0]:
                        pref_match = re.search(r'#(\d)', cells[0])
                        if pref_match:
                            pref_rank = int(pref_match.group(1))
                            pref_data = {'preference_rank': pref_rank}

                            # Look ahead for the next few rows to get data
                            pref_data.update(self._extract_preference_from_table(table, i, pref_rank))

                            if self._is_valid_preference(pref_data):
                                self.data['preferences'].append(pref_data)

    def _extract_preference_from_table(self, table, start_row, pref_rank):
        """Extract a single preference's data from table rows."""
        pref_data = {'preference_rank': pref_rank}  # FIX: Add preference_rank

        # Look at the next few rows after the preference header
        for i in range(start_row, min(start_row + 5, len(table.rows))):
            row = table.rows[i]
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) >= 2:
                field_name = cells[0].lower()
                field_value = cells[1]

                if 'hut name' in field_name:
                    # Check for officially sanctioned designation
                    if 'officially sanctioned' in field_value.lower():
                        pref_data['sanctioned'] = True
                    pref_data['hut_name'] = field_value
                elif 'date in' in field_name:
                    pref_data['date_in'] = self._parse_date(field_value)
                elif 'date out' in field_name:
                    pref_data['date_out'] = self._parse_date(field_value)
                elif 'number of guests' in field_name or 'guests' in field_name:
                    # FIX: Better handling of ENTIRE variants (En/re, Entire, ENTIRE, etc.)
                    field_value_upper = field_value.upper()
                    # Check for various spellings: ENTIRE, EN/RE, ENTTRE, etc.
                    is_entire = (
                        'ENTIRE' in field_value_upper or
                        'EN/RE' in field_value_upper or
                        'ENTRE' in field_value_upper or
                        'ENTTRE' in field_value_upper or
                        re.search(r'ENT[/\\]?[IT]?RE', field_value_upper)
                    )
                    if is_entire:
                        pref_data['party_size'] = 'ENTIRE'
                    else:
                        # FIX: Only try to extract number if digits exist
                        digit_match = re.search(r'\d+', field_value)
                        if digit_match:
                            pref_data['party_size'] = int(digit_match.group())

        return pref_data

    def _extract_from_paragraphs(self):
        """Extract data from document paragraphs (fallback method)."""
        text = '\n'.join([para.text for para in self.doc.paragraphs])

        # Use similar logic to PDF parser
        lines = text.split('\n')

        # Extract leader info
        for i, line in enumerate(lines):
            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', line)
            if email_match:
                self.data['email'] = email_match.group()

            phone_match = re.search(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', line)
            if phone_match:
                self.data['phone'] = phone_match.group()

            if 'Leader Name' in line and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and 'Email' not in next_line:
                    self.data['leader_name'] = next_line

        # Extract preferences
        current_pref = None
        pref_data = {}

        for i, line in enumerate(lines):
            line = line.strip()

            pref_match = re.match(r'Hut Preference #(\d)', line)
            if pref_match:
                if current_pref and self._is_valid_preference(pref_data):
                    self.data['preferences'].append(pref_data)

                current_pref = int(pref_match.group(1))
                pref_data = {'preference_rank': current_pref}
                continue

            if current_pref is None:
                continue

            if 'Hut Name' in line and i + 1 < len(lines):
                hut_name = lines[i + 1].strip()
                if hut_name and 'Date' not in hut_name:
                    pref_data['hut_name'] = hut_name

            elif 'Date In' in line and i + 1 < len(lines):
                date_in = lines[i + 1].strip()
                if date_in and self._looks_like_date(date_in):
                    pref_data['date_in'] = self._parse_date(date_in)

            elif 'Date Out' in line and i + 1 < len(lines):
                date_out = lines[i + 1].strip()
                if date_out and self._looks_like_date(date_out):
                    pref_data['date_out'] = self._parse_date(date_out)

            elif 'Number of Guests' in line and i + 1 < len(lines):
                guests = lines[i + 1].strip()
                if guests and guests.upper() != 'ENTIRE':
                    try:
                        pref_data['party_size'] = int(re.search(r'\d+', guests).group())
                    except:
                        pass
                elif guests.upper() == 'ENTIRE':
                    pref_data['party_size'] = 'ENTIRE'

        if current_pref and self._is_valid_preference(pref_data):
            self.data['preferences'].append(pref_data)

    def _looks_like_date(self, text):
        """Check if text looks like a date."""
        patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
            r'[A-Za-z]{3,}\s+\d{1,2},?\s+\d{4}',
        ]
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        return False

    def _parse_date(self, date_str):
        """Parse various date formats."""
        date_str = date_str.strip()

        formats = [
            '%m/%d/%Y',
            '%m-%d-%Y',
            '%Y-%m-%d',
            '%m/%d/%y',
            '%m-%d-%y',
            '%B %d, %Y',
            '%b %d, %Y',
            '%B %d %Y',
            '%b %d %Y',
        ]

        for fmt in formats:
            try:
                return datetime.strptime(date_str, fmt).strftime('%Y-%m-%d')
            except:
                continue

        # FIX: Handle M/D format without year (assume 2026)
        # Try to parse M/D and add year
        short_formats = [
            '%m/%d',
            '%m-%d',
        ]

        for fmt in short_formats:
            try:
                dt = datetime.strptime(date_str, fmt)
                # Assume 2026 for dates without year
                dt = dt.replace(year=2026)
                return dt.strftime('%Y-%m-%d')
            except:
                continue

        return date_str

    def _is_valid_preference(self, pref_data):
        """Check if preference has minimum required data."""
        required = ['preference_rank', 'hut_name', 'date_in', 'date_out', 'party_size']
        return all(key in pref_data for key in required)

    def to_csv_rows(self):
        """Convert parsed data to CSV row format."""
        rows = []
        leader_name = self.data.get('leader_name', 'Unknown')

        for pref in self.data['preferences']:
            rows.append({
                'UserName': leader_name,
                'PreferenceRank': pref['preference_rank'],
                'Hut': pref['hut_name'],
                'StartDate': pref['date_in'],
                'EndDate': pref['date_out'],
                'PartySize': pref['party_size'],
                'Sanctioned': 'YES' if pref.get('sanctioned', False) else ''
            })

        return rows
